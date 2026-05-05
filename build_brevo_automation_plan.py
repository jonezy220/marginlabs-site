#!/usr/bin/env python3
"""
Margin Labs — Brevo Marketing Automation Plan (Word Document Builder)
Run with: python3 build_brevo_automation_plan.py
Output:   Brevo_Marketing_Automation_Plan_v1.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_FILE = "Brevo_Marketing_Automation_Plan_v1.docx"

NAVY   = RGBColor(0x1F, 0x38, 0x64)
COPPER = RGBColor(0xC8, 0x82, 0x3C)
IRON   = RGBColor(0x22, 0x22, 0x22)
GRAY   = RGBColor(0x59, 0x59, 0x59)
BLACK  = RGBColor(0x00, 0x00, 0x00)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LTGRAY = RGBColor(0xF5, 0xF5, 0xF0)
GREEN  = RGBColor(0x37, 0x86, 0x4A)
RED    = RGBColor(0xC0, 0x39, 0x2B)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CCCCCC")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def setup_styles(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = IRON
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = Pt(15)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(26)
    run.font.color.rgb = NAVY
    run.bold = True


def subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = COPPER
    run.bold = False


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY
    run.bold = True
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "C8823C")
    pBdr.append(bottom)
    pPr.append(pBdr)


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = COPPER
    run.bold = True


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    run.bold = True


def body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = color or IRON
    run.bold = bold
    run.italic = italic
    return p


def bullet(doc, text, level=0, bold=False):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = IRON
    run.bold = bold


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.3)
    # Light background via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFF8F0")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = GRAY
    run.italic = True


def email_block(doc, subject, body_text, from_addr="chris@marginlabs.io", day="", condition=""):
    """Render a styled email preview block."""
    # Email header bar
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "1F3864")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    if day:
        run = p.add_run(f"Day {day}  |  ")
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.bold = False
    run = p.add_run(f"SUBJECT: {subject}")
    run.font.size = Pt(10)
    run.font.color.rgb = WHITE
    run.bold = True

    # Email body
    tbl2 = doc.add_table(rows=1, cols=1)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell2 = tbl2.rows[0].cells[0]
    set_cell_bg(cell2, "F9F9F6")
    # Add border
    tc = cell2._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["left", "right", "bottom"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "DDDDDD")
        tcBorders.append(b)
    tcPr.append(tcBorders)

    p = cell2.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"FROM: {from_addr}\n\n")
    run.font.size = Pt(8.5)
    run.font.color.rgb = GRAY
    run.font.name = "Consolas"

    for line in body_text.strip().split("\n"):
        run = p.add_run(line + "\n")
        if line.startswith("→"):
            run.font.size = Pt(10)
            run.font.color.rgb = COPPER
            run.bold = True
        elif line.startswith("──") or line.startswith("────"):
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        elif line.startswith("  ·") or line.startswith("  "):
            run.font.size = Pt(10)
            run.font.color.rgb = IRON
        else:
            run.font.size = Pt(10)
            run.font.color.rgb = IRON

    if condition:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(8)
        run = p2.add_run(f"BREVO CONDITION: {condition}")
        run.font.size = Pt(9)
        run.font.color.rgb = COPPER
        run.bold = True
    else:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(8)


def add_data_table(doc, headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(tbl)

    # Header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, "1F3864")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(h)
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.bold = True

    # Data rows
    for r_idx, row in enumerate(rows):
        bg = "FFFFFF" if r_idx % 2 == 0 else "F9F9F6"
        for c_idx, val in enumerate(row):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = IRON

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def build(doc):
    # ── COVER ──
    title(doc, "Brevo Marketing Automation Plan")
    subtitle(doc, "Version 1.0  ·  April 2026  ·  Based on Business Plan v8")

    body(doc, "All automation runs in Brevo. Resend handles transactional emails only "
         "(purchase confirmations, PDF delivery). Everything below is configured as "
         "Brevo Automation Workflows.", italic=True, color=GRAY)

    body(doc, "14 emails across 4 workflows. 4 funnel entry points. "
         "Designed to run unattended once configured.", bold=True)

    # ── FUNNEL OVERVIEW ──
    h1(doc, "Funnel Overview")

    body(doc, "All traffic enters through one of three gates. Each gate feeds a "
         "dedicated nurture workflow. The two lead-nurture workflows (Free Guide "
         "and Margin Multiplier) cross-sell each other and converge on the $139 "
         "Framework purchase. Post-purchase nurture drives the $379 Quick Start "
         "Call as the primary conversion target.")

    add_data_table(doc,
        ["Entry Point", "Gate", "Workflow", "Emails", "Goal"],
        [
            ["Free Guide", "Email capture", "WF 1: Free Guide Nurture", "4 (Days 2-14)", "Run Multiplier → Buy Framework"],
            ["Margin Multiplier", "Email capture", "WF 2: Multiplier Nurture", "4 (Days 2-14)", "Get Free Guide → Buy Framework"],
            ["Framework Purchase", "Stripe checkout", "WF 3: Post-Purchase", "4 (Days 3-21)", "Book Quick Start Call ($379)"],
            ["Advisory Form", "Consult form", "WF 4: Advisory Confirm", "2 (Day 0, 3)", "Confirm receipt → manual follow-up"],
        ]
    )

    # Funnel flow
    h2(doc, "Funnel Flow")
    body(doc, "TRAFFIC  →  Free Guide / Multiplier / Advisory Form")
    body(doc, "           ↓                    ↓                    ↓")
    body(doc, "     WF1 Nurture      WF2 Nurture         WF4 Confirmation")
    body(doc, "     (cross-sell         (cross-sell")
    body(doc, "      Multiplier)        Free Guide)")
    body(doc, "           ↓                    ↓")
    body(doc, "           └────────┬───────────┘")
    body(doc, "                    ↓")
    body(doc, "           $139 FRAMEWORK PURCHASE", bold=True)
    body(doc, "                    ↓")
    body(doc, "             WF3 Post-Purchase")
    body(doc, "                    ↓")
    body(doc, "     $379 QUICK START CALL  (primary CTA)", bold=True, color=COPPER)
    body(doc, "         or  $697 PLAYBOOK  (DIY secondary)", italic=True)

    # ── BREVO CONFIGURATION ──
    add_page_break(doc)
    h1(doc, "Brevo Configuration")

    h2(doc, "Contact Attributes")
    body(doc, "Configure in Brevo → Contacts → Settings → Contact Attributes.")

    add_data_table(doc,
        ["Attribute", "Type", "Values", "Purpose"],
        [
            ["SOURCE", "Text", "Free Guide / Margin Multiplier / Advisory Form / Framework Customer", "Already exists. Set by API on entry."],
            ["HAS_FREE_GUIDE", "Boolean", "true / false", "Tracks whether contact has the free guide"],
            ["HAS_MULTIPLIER", "Boolean", "true / false", "Tracks whether contact ran the calculator"],
            ["HAS_FRAMEWORK", "Boolean", "true / false", "Set true on $139 purchase"],
            ["ENTRY_DATE", "Date", "Auto", "When contact first entered the system"],
            ["PAYMENTS_VOLUME", "Text", "(existing)", "Monthly payment volume from calculator"],
            ["RECOMMENDED_MODEL", "Text", "(existing)", "Model recommended by Multiplier"],
            ["OPPORTUNITY_GAP", "Text", "(existing)", "Revenue gap identified by Multiplier"],
        ]
    )

    h2(doc, "Tags (Workflow Triggers)")
    body(doc, "Tags are applied by API functions when contacts take actions on the site.")

    add_data_table(doc,
        ["Tag", "Applied By", "Purpose"],
        [
            ["free-guide-lead", "api/send-guide.js", "Entered via free guide download"],
            ["multiplier-lead", "api/submit-calculator.js", "Entered via Margin Multiplier"],
            ["advisory-lead", "api/submit-consult.js", "Submitted advisory/consult form"],
            ["framework-customer", "api/stripe-webhook.js", "Purchased $139 Framework"],
            ["playbook-customer", "api/stripe-webhook.js (future)", "Purchased $697 Playbook"],
            ["call-booked", "Manual", "Booked Quick Start Call"],
        ]
    )

    h2(doc, "Suppression & Conflict Rules")

    add_data_table(doc,
        ["Event", "Action"],
        [
            ["Contact enters WF1 AND already has multiplier-lead tag", "Skip WF1 entirely — they follow WF2"],
            ["Contact enters WF2 AND already in WF1", "Exit WF1 immediately; WF2 takes over"],
            ["Contact gets framework-customer tag", "Exit WF1 and WF2 immediately; enter WF3"],
            ["Contact gets playbook-customer tag", "Exit WF3 immediately"],
            ["Contact gets call-booked tag", "Exit WF3 immediately"],
            ["Contact unsubscribes", "Exit all workflows (Brevo handles natively)"],
        ]
    )

    note(doc, "Merge rule: If a contact has BOTH free-guide-lead and multiplier-lead, they follow "
         "Workflow 2 only. Multiplier results are more personalized and create a stronger hook. "
         "Workflow 1 exits via suppression.")

    # ── WORKFLOW 1 ──
    add_page_break(doc)
    h1(doc, "Workflow 1: Free Guide Nurture")

    body(doc, "Trigger: Tag free-guide-lead added", bold=True)
    body(doc, "Suppression: Skip if tag framework-customer exists", bold=True)
    body(doc, "Goal: Drive Margin Multiplier usage → Framework purchase", bold=True)

    h3(doc, "Email 1.1 — Day 2")
    email_block(doc,
        subject="A tool that goes deeper than the guide",
        day="2",
        body_text="""MARGIN LABS

Hope you've had a chance to look through the guide.

It covers the landscape — the four models, what each one means, and the questions to ask before you start evaluating. It's the orientation layer.

What it can't do is tell you what the numbers look like for YOUR platform.

That's what the Margin Multiplier does. You plug in your payment volume and it runs a side-by-side comparison across all four monetization models — ISV Referral, Enhanced Residuals, PayFac-as-a-Service, and Full PayFac.

Takes about 60 seconds:
→ marginlabs.io/margin-multiplier

The output pairs well with Section 2 of the guide (the model comparison). Worth doing them together.

— Chris, Margin Labs""",
        condition="If HAS_MULTIPLIER = true, skip Email 1.2 and go to 1.3."
    )

    h3(doc, "Email 1.2 — Day 5")
    email_block(doc,
        subject="The number most platforms don't know",
        day="5",
        body_text="""MARGIN LABS

There's one number that determines whether embedded payments is a line item or a business unit for your platform.

It's your opportunity gap — the delta between what you're earning today on payments and what you'd earn under the optimal model at your volume.

Most platforms don't know this number. Their processor doesn't volunteer it.

The Margin Multiplier calculates it in about 60 seconds.

→ marginlabs.io/margin-multiplier

If the gap is under $50K/year, you're probably fine where you are. If it's over $200K, the conversation gets interesting.

— Chris"""
    )

    h3(doc, "Email 1.3 — Day 9")
    email_block(doc,
        subject="What the guide doesn't cover (deliberately)",
        day="9",
        body_text="""MARGIN LABS

The free guide gives you the landscape. It doesn't give you:

  · The actual take rate economics at your volume level
  · Which vendors lead in each model category
  · How to assess your platform's readiness for each model
  · A methodology for building a business case internally
  · An implementation roadmap with phase-by-phase milestones

Those gaps are deliberate. The guide is the "should we think about this?" layer. The Strategic Framework is the "how do we think about this?" layer.

45 pages. Covers model economics, vendor landscape, readiness assessment, ROI methodology, and implementation planning.

→ marginlabs.io/strategic-framework ($139)

— Chris"""
    )

    h3(doc, "Email 1.4 — Day 14")
    email_block(doc,
        subject="Last note from me on this",
        day="14",
        body_text="""MARGIN LABS

Final email in this sequence.

If embedded payments is on your roadmap — even loosely — the Strategic Framework ($139) is the most efficient way to get a real picture of the economics and what it takes.

If the timing isn't right, no problem. The Lab has free articles on the topics that come up most:

→ marginlabs.io/the-lab

I'll only reach out again when there's something worth sending.

— Chris"""
    )

    note(doc, "End of Workflow 1. Contact stays in Brevo list for future broadcasts.")

    # ── WORKFLOW 2 ──
    add_page_break(doc)
    h1(doc, "Workflow 2: Margin Multiplier Nurture")

    body(doc, "Trigger: Tag multiplier-lead added", bold=True)
    body(doc, "Suppression: Skip if tag framework-customer exists", bold=True)
    body(doc, "Goal: Deliver free guide (if they don't have it) → Framework purchase", bold=True)

    h3(doc, "Email 2.1 — Day 2")
    email_block(doc,
        subject="Your results + a resource that pairs with them",
        day="2",
        body_text="""MARGIN LABS

A couple of days since you ran the Margin Multiplier.

Your results showed {{RECOMMENDED_MODEL}} as the highest-margin model at your volume. The opportunity gap — what you'd gain by moving to that model — was approximately {{OPPORTUNITY_GAP}}.

That's a directional estimate. The actual number depends on your merchant mix, vertical, and the terms you negotiate.

If you haven't grabbed it yet, the free guide gives you the landscape view — what each model actually means, the questions to ask, and the common mistakes platforms make early on.

→ Download the free guide: marginlabs.io
  (scroll to "Get the Free Guide" — just your email)

It pairs well with the Multiplier output. The guide explains the models; the Multiplier shows your specific numbers.

— Chris""",
        condition="If HAS_FREE_GUIDE = true, replace free guide offer with link to Lab article (Why Merchants Don't Use Payments)."
    )

    h3(doc, "Email 2.2 — Day 5")
    email_block(doc,
        subject="The activation problem nobody talks about",
        day="5",
        body_text="""MARGIN LABS

Your Multiplier results assume 100% merchant activation. In practice, most platforms see 40-70%.

That means the real number is probably lower than the estimate you saw — unless you solve the activation problem.

A platform with 80% activation on Enhanced Residuals generates more revenue than a platform with 40% activation on PayFac-as-a-Service. The model matters, but activation matters more.

More on this:
→ marginlabs.io/the-lab/why-merchants-dont-use-payments

That article covers the three reasons merchants don't opt in and what the top-performing platforms do differently.

— Chris"""
    )

    h3(doc, "Email 2.3 — Day 9")
    email_block(doc,
        subject="What the Multiplier doesn't show you",
        day="9",
        body_text="""MARGIN LABS

The Margin Multiplier gives you the headline number — the revenue estimate across all four models at your volume.

What it doesn't show:

  · Why the take rate ranges vary so much within each model
  · Which vendors are strongest at your volume tier
  · What it actually costs to implement each model
  · Whether your platform is ready (technically and operationally)
  · How to build a business case that survives a board meeting

That's what the Strategic Framework covers. 45 pages. Model economics, vendor landscape, readiness assessment, ROI methodology, and implementation planning.

The Multiplier shows you the opportunity exists. The Framework shows you how to capture it.

→ marginlabs.io/strategic-framework ($139)

— Chris"""
    )

    h3(doc, "Email 2.4 — Day 14")
    email_block(doc,
        subject="Last note on this",
        day="14",
        body_text="""MARGIN LABS

Final email in this sequence.

Your Multiplier estimate showed a {{OPPORTUNITY_GAP}} opportunity gap. If that number got your attention, the Strategic Framework ($139) is the next step — it turns that estimate into a plan you can act on.

If the timing isn't right:
→ marginlabs.io/the-lab
has free content on the topics that matter most.

I'll only reach out again when there's something worth sending.

— Chris"""
    )

    note(doc, "End of Workflow 2. Contact stays in Brevo list for future broadcasts.")

    # ── WORKFLOW 3 ──
    add_page_break(doc)
    h1(doc, "Workflow 3: Post-Purchase Nurture ($139 Framework)")

    body(doc, "Trigger: Tag framework-customer added", bold=True)
    body(doc, "Suppression: Immediately remove contact from Workflows 1 and 2", bold=True)
    body(doc, "Goal: Drive Quick Start Call ($379) booking. Playbook ($697) as secondary DIY option.", bold=True)

    h3(doc, "Email 3.1 — Day 3")
    email_block(doc,
        subject="Getting the most out of the Framework",
        day="3",
        body_text="""MARGIN LABS

Your Framework should have arrived. A few things worth reading closely.

Section 2 (Model Economics): The take rate ranges are real-world benchmarks, not theoretical. If your processor is quoting outside those ranges, that's useful information for your next conversation.

Section 4 (Readiness Assessment): This is the section most operators skip and then wish they hadn't. The model decision is only as good as your platform's ability to execute it.

Section 5 (ROI Methodology): Use this to build your own projection with your own numbers. The Multiplier gave you a directional estimate — this section shows you how to make it defensible.

Questions on anything in there? Reply to this email. I read every one.

— Chris, Margin Labs"""
    )

    h3(doc, "Email 3.2 — Day 7")
    email_block(doc,
        subject="The fastest way to go from framework to decision",
        day="7",
        body_text="""MARGIN LABS

By now you've worked through at least part of the Framework.

Most operators land in one of two places:

  A) "I know which model makes sense. Now I need to figure out the vendor, the terms, and the implementation plan."

  B) "I have better questions than when I started, but I'm not sure how to apply this to my specific situation."

Either way — the fastest path forward is a conversation, not more reading.

────────────────────────────────────────
QUICK START CALL — $379

One hour. Your specific situation, your numbers, your platform. We'll cover:

  · Which model fits your volume and stage
  · What realistic economics look like for you specifically
  · Which vendors to talk to (and which to skip)
  · What to ask for in your first processor conversation
  · Warm introductions to vetted processing partners

You'll get a follow-up email with specific action items and next steps within 48 hours.

→ marginlabs.io/advisory
────────────────────────────────────────

If you'd rather go the DIY route, the Execution Playbook ($697) has vendor scorecards, contract term benchmarks, negotiation playbook, and implementation project plans.

→ marginlabs.io/execution-playbook

The call gets you to a decision faster. The Playbook gets you there on your own timeline.

— Chris"""
    )

    h3(doc, "Email 3.3 — Day 14")
    email_block(doc,
        subject="Where does this land for you?",
        day="14",
        body_text="""MARGIN LABS

Two weeks since you picked up the Framework.

Curious where things stand. The operators I work with usually fall into a pattern:

  Week 1: Read through, run the numbers, identify the model
  Week 2: Start thinking about vendors and internal buy-in
  Week 3-4: Either start vendor conversations or realize they need help navigating the specifics

If you're at the vendor conversation stage — or heading there — the Quick Start Call is designed for exactly that moment. One hour, your situation, specific recommendations and introductions.

→ marginlabs.io/advisory ($379)

If you're not there yet, no pressure. Reply here if anything in the Framework needs clarification.

— Chris"""
    )

    h3(doc, "Email 3.4 — Day 21")
    email_block(doc,
        subject="One last thought",
        day="21",
        body_text="""MARGIN LABS

Last note from me.

Two paths forward from the Framework:

  GUIDED: Quick Start Call ($379)
  One hour on your specific situation. Model recommendation, vendor shortlist, negotiation guidance, warm intros to vetted processing partners. Follow-up action plan included.
→ marginlabs.io/advisory

  DIY: Execution Playbook ($697)
  Vendor scorecards, contract benchmarks, negotiation playbook, ROI calculator, implementation project plans, board presentation template. Everything you need to execute independently.
→ marginlabs.io/execution-playbook

Either path works. The call is faster and includes introductions. The Playbook is more comprehensive and works on your own timeline.

After this, I'll only be in touch when I have something genuinely useful to send.

— Chris"""
    )

    note(doc, "End of Workflow 3.")

    # ── WORKFLOW 4 ──
    add_page_break(doc)
    h1(doc, "Workflow 4: Advisory Confirmation")

    body(doc, "Trigger: Tag advisory-lead added", bold=True)
    body(doc, "Suppression: None (advisory leads can also be in product workflows)", bold=True)
    body(doc, "Goal: Confirm receipt. One follow-up. Chris handles manually from here.", bold=True)

    h3(doc, "Email 4.1 — Immediate")
    email_block(doc,
        subject="Got your inquiry — we'll be in touch",
        day="0",
        body_text="""MARGIN LABS

Thanks for reaching out. We've received your information and will follow up within 48 hours.

In the meantime — if you haven't run the Margin Multiplier, it'll give us a cleaner starting point for the conversation:

→ marginlabs.io/margin-multiplier

— Chris, Margin Labs"""
    )

    h3(doc, "Email 4.2 — Day 3")
    email_block(doc,
        subject="Quick follow-up on your inquiry",
        day="3",
        body_text="""MARGIN LABS

Following up on your inquiry from a few days ago. I want to make sure my response didn't land in spam — I'll be reaching out directly shortly.

If anything has changed or you want to add context before we connect, reply here.

— Chris"""
    )

    note(doc, "End of Workflow 4. Everything after this is manual outreach by Chris.")

    # ── SEQUENCE TIMING SUMMARY ──
    add_page_break(doc)
    h1(doc, "Sequence Timing Summary")

    add_data_table(doc,
        ["Workflow", "Email", "Day", "Subject", "CTA"],
        [
            ["1 Free Guide",    "1.1", "2",  "A tool that goes deeper than the guide",              "Run Multiplier"],
            ["1 Free Guide",    "1.2", "5",  "The number most platforms don't know",                "Run Multiplier"],
            ["1 Free Guide",    "1.3", "9",  "What the guide doesn't cover (deliberately)",         "Buy Framework $139"],
            ["1 Free Guide",    "1.4", "14", "Last note from me on this",                           "Buy Framework $139"],
            ["2 Multiplier",    "2.1", "2",  "Your results + a resource that pairs with them",      "Get Free Guide"],
            ["2 Multiplier",    "2.2", "5",  "The activation problem nobody talks about",           "Lab article"],
            ["2 Multiplier",    "2.3", "9",  "What the Multiplier doesn't show you",                "Buy Framework $139"],
            ["2 Multiplier",    "2.4", "14", "Last note on this",                                   "Buy Framework $139"],
            ["3 Post-Purchase", "3.1", "3",  "Getting the most out of the Framework",               "Engage with product"],
            ["3 Post-Purchase", "3.2", "7",  "The fastest way to go from framework to decision",    "Book QSC $379"],
            ["3 Post-Purchase", "3.3", "14", "Where does this land for you?",                       "Book QSC $379"],
            ["3 Post-Purchase", "3.4", "21", "One last thought",                                    "QSC $379 / Playbook $697"],
            ["4 Advisory",      "4.1", "0",  "Got your inquiry — we'll be in touch",                "Run Multiplier"],
            ["4 Advisory",      "4.2", "3",  "Quick follow-up on your inquiry",                     "Reply"],
        ]
    )

    # ── API UPDATES ──
    h1(doc, "API Updates Required")

    body(doc, "These changes ensure the site correctly feeds contact data and tags to Brevo, "
         "triggering the right workflows automatically.")

    add_data_table(doc,
        ["API File", "Change Required"],
        [
            ["api/send-guide.js",          "Add tag free-guide-lead; set HAS_FREE_GUIDE = true"],
            ["api/submit-calculator.js",   "Add tag multiplier-lead; set HAS_MULTIPLIER = true"],
            ["api/submit-consult.js",      "Add tag advisory-lead"],
            ["api/stripe-webhook.js",      "Add tag framework-customer; set HAS_FRAMEWORK = true; set SOURCE = Framework Customer"],
            ["api/brevo-subscribe.js",     "Verify OPPORTUNITY_GAP and RECOMMENDED_MODEL are being passed (already partially wired)"],
        ]
    )

    # ── IMPLEMENTATION CHECKLIST ──
    h1(doc, "Implementation Checklist")

    h2(doc, "Brevo Configuration")
    bullet(doc, "Add contact attributes: HAS_FREE_GUIDE, HAS_MULTIPLIER, HAS_FRAMEWORK, ENTRY_DATE")
    bullet(doc, "Build email template matching site aesthetic (dark #111111, copper #C8823C, DM Sans/DM Mono, text-first)")
    bullet(doc, "Create Workflow 1 (Free Guide Nurture) — 4 emails, Days 2/5/9/14")
    bullet(doc, "Create Workflow 2 (Multiplier Nurture) — 4 emails, Days 2/5/9/14")
    bullet(doc, "Create Workflow 3 (Post-Purchase) — 4 emails, Days 3/7/14/21")
    bullet(doc, "Create Workflow 4 (Advisory Confirmation) — 2 emails, Immediate/Day 3")
    bullet(doc, "Configure suppression rules between workflows")
    bullet(doc, "Test with chris@marginlabs.io through each entry point")

    h2(doc, "API Updates")
    bullet(doc, "api/send-guide.js — add tag free-guide-lead, set HAS_FREE_GUIDE = true")
    bullet(doc, "api/submit-calculator.js — add tag multiplier-lead, set HAS_MULTIPLIER = true")
    bullet(doc, "api/submit-consult.js — add tag advisory-lead")
    bullet(doc, "api/stripe-webhook.js — add tag framework-customer, set HAS_FRAMEWORK = true")

    # ── WHAT THIS REPLACES ──
    h1(doc, "What This Replaces")

    body(doc, "This plan supersedes two older documents:")
    bullet(doc, "brevo-sequences.md — 5 sequences, less structured, partially overlapping")
    bullet(doc, "email-waterfalls-plan.md — 6 waterfalls, based on v5 pricing, proposed Resend+Postgres stack instead of Brevo")
    body(doc, "Both older docs can be archived once this plan is approved and built.")

    # ── FUTURE ──
    h1(doc, "Future Enhancements (Not in v1)")

    bullet(doc, "Volume-based segmentation (High-Value / Mid-Market / Early-Stage tracks)")
    bullet(doc, "Progressive profiling survey (Day 1 email with segment self-selection)")
    bullet(doc, "Post-Playbook purchase sequence (→ Quick Start Call upsell)")
    bullet(doc, "Post-Quick Start Call follow-up (Day 30 check-in)")
    bullet(doc, "Quarterly re-engagement for dormant leads")
    bullet(doc, "A/B testing subject lines in Brevo")


def main():
    doc = Document()
    setup_styles(doc)
    build(doc)
    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), OUTPUT_FILE)
    doc.save(out_path)
    print(f"Built: {out_path}")


if __name__ == "__main__":
    main()
